#!/usr/bin/env python3
"""
File utilities for letter of recommendation tools.

Handles file operations like finding documents, saving files, etc.
"""

import logging
from pathlib import Path
from typing import List, Dict, Optional
import mammoth

logger = logging.getLogger(__name__)

def convert_docx_to_markdown(docx_path: Path) -> str:
    """Convert a Word document to Markdown format, dropping all images."""
    try:
        logger.info(f"Converting {docx_path.name} to Markdown...")
        with open(docx_path, 'rb') as docx_file:
            # Configure mammoth to ignore images by returning empty strings
            result = mammoth.convert_to_markdown(
                docx_file,
                convert_image=mammoth.images.inline(lambda image: {"alt": ""})
            )
            markdown_text = result.value

            # Log any conversion warnings
            if result.messages:
                for message in result.messages:
                    logger.debug(f"Conversion message: {message}")

            logger.info(f"Successfully converted {docx_path.name} to Markdown (images dropped)")
            return markdown_text

    except Exception as e:
        logger.error(f"Error converting {docx_path}: {e}")
        raise

def find_docx_files(path: Path) -> List[Path]:
    """Find all .docx files in a path (file or directory)."""
    if path.is_file():
        if path.suffix.lower() == '.docx' and not path.name.startswith('~$'):
            return [path]
        else:
            logger.warning(f"File {path} is not a .docx file")
            return []
    elif path.is_dir():
        # Recursively find all .docx files, excluding temporary files
        docx_files = [f for f in path.rglob('*.docx') if not f.name.startswith('~$')]
        logger.info(f"Found {len(docx_files)} .docx file(s) in {path}")
        return docx_files
    else:
        logger.error(f"Path {path} does not exist")
        return []


def convert_pdf_to_markdown(pdf_path: Path) -> str:
    """
    Convert a PDF to Markdown format using PyPDF2.

    Note: This is a simple text extraction. For better results with complex PDFs,
    consider using more advanced tools like pdfplumber or pdf2image + OCR.
    """
    try:
        import PyPDF2
        logger.info(f"Converting {pdf_path.name} to Markdown...")

        with open(pdf_path, 'rb') as pdf_file:
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            text_parts = []

            for page_num, page in enumerate(pdf_reader.pages, 1):
                text = page.extract_text()
                if text.strip():
                    text_parts.append(text)
                else:
                    logger.warning(f"Page {page_num} of {pdf_path.name} appears empty or unreadable")

            markdown_text = "\n\n".join(text_parts)
            logger.info(f"Successfully converted {pdf_path.name} to Markdown ({len(pdf_reader.pages)} pages)")
            return markdown_text

    except ImportError:
        logger.error("PyPDF2 not installed. Install it with: pip install PyPDF2")
        raise
    except Exception as e:
        logger.error(f"Error converting PDF {pdf_path}: {e}")
        raise


def convert_file_to_markdown(file_path: Path) -> str:
    """
    Convert a file to Markdown format based on its extension.

    Supports: .docx, .pdf, .txt, .md
    """
    suffix = file_path.suffix.lower()

    if suffix == '.docx':
        return convert_docx_to_markdown(file_path)
    elif suffix == '.pdf':
        return convert_pdf_to_markdown(file_path)
    elif suffix in ['.txt', '.md']:
        logger.info(f"Reading text file {file_path.name}...")
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    else:
        raise ValueError(f"Unsupported file type: {suffix}. Supported types: .docx, .pdf, .txt, .md")


def find_student_materials(input_dir: Path) -> Dict[str, Path]:
    """
    Find student materials in the input directory.

    Looks for common filenames and returns a dict mapping material type to file path.
    Returns empty dict if directory doesn't exist.
    """
    if not input_dir.exists():
        logger.warning(f"Input directory does not exist: {input_dir}")
        return {}

    materials = {}

    # Common patterns for each material type
    patterns = {
        'resume': ['resume.*', 'cv.*'],
        'transcript': ['transcript.*', 'grades.*'],
        'accomplishments': ['accomplishments.*', 'achievements.*', 'activities.*'],
        'statement': ['statement.*', 'personal_statement.*', 'sop.*', 'purpose.*']
    }

    for material_type, pattern_list in patterns.items():
        for pattern in pattern_list:
            matches = list(input_dir.glob(pattern))
            if matches:
                # Take the first match for each type
                materials[material_type] = matches[0]
                logger.info(f"Found {material_type}: {matches[0].name}")
                break

    return materials


def convert_markdown_to_docx(markdown_path: Path, docx_path: Optional[Path] = None) -> Path:
    """
    Convert a Markdown file to DOCX format.

    Args:
        markdown_path: Path to input .md file
        docx_path: Path for output .docx file (optional, defaults to same name with .docx extension)

    Returns:
        Path to created DOCX file
    """
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        logger.info(f"Converting {markdown_path.name} to DOCX...")

        # Set default output path
        if docx_path is None:
            docx_path = markdown_path.with_suffix('.docx')

        # Read markdown content
        with open(markdown_path, 'r', encoding='utf-8') as f:
            content = f.read()

        # Create document
        doc = Document()

        # Set default font
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)

        # Process content line by line
        lines = content.split('\n')
        i = 0
        while i < len(lines):
            line = lines[i].strip()

            # Skip empty lines at start
            if not line:
                i += 1
                continue

            # Handle letterhead (lines starting with __)
            if line.startswith('__') and line.endswith('__'):
                text = line.strip('_')
                p = doc.add_paragraph(text)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = p.runs[0]
                run.bold = True
                i += 1
                continue

            # Regular paragraph
            if line and not line.startswith('#'):
                doc.add_paragraph(line)

            i += 1

        # Save document
        docx_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(docx_path)

        logger.info(f"Successfully converted to: {docx_path}")
        return docx_path

    except ImportError:
        logger.error("python-docx not installed. This is required for DOCX conversion.")
        raise
    except Exception as e:
        logger.error(f"Error converting to DOCX: {e}")
        raise


def save_markdown(content: str, output_path: Path) -> None:
    """Save markdown content to a file."""
    try:
        output_path.parent.mkdir(parents=True, exist_ok=True)
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(content)
        logger.info(f"Saved content to {output_path}")

    except Exception as e:
        logger.error(f"Error saving file {output_path}: {e}")
        raise
