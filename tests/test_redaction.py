#!/usr/bin/env python3
"""
System test for the redaction script.
Creates a sample .docx file with student information, runs the redaction script,
and verifies the output is properly redacted.
"""

import os
import subprocess
import sys
import tempfile
from pathlib import Path

from docx import Document


def create_test_document(output_path: Path) -> dict:
    """
    Create a test Word document with known student information.
    Returns a dict of student info that should be redacted.
    """
    doc = Document()

    # Add title
    doc.add_heading('Letter of Recommendation', 0)

    # Add content with student information
    student_info = {
        'name': 'John Smith',
        'email': 'jsmith@university.edu',
        'student_id': 'STU123456'
    }

    doc.add_paragraph(
        f"To Whom It May Concern,\n\n"
        f"I am writing to recommend {student_info['name']} for your graduate program. "
        f"{student_info['name']} has been an outstanding student in my Computer Science course.\n\n"
        f"You can reach {student_info['name']} at {student_info['email']} or reference "
        f"student ID {student_info['student_id']} in your records.\n\n"
        f"{student_info['name']} has demonstrated excellent problem-solving skills and "
        f"a deep understanding of algorithms and data structures.\n\n"
        f"Sincerely,\n"
        f"Professor Jane Doe\n"
        f"Department of Computer Science\n"
        f"University of Example"
    )

    doc.save(str(output_path))
    print(f"✓ Created test document: {output_path}")
    return student_info


def run_redaction_script(docx_path: Path) -> int:
    """
    Run the redaction script on the test document.
    Returns the exit code.
    """
    script_path = Path(__file__).parent.parent / 'redact_student_info.py'

    print(f"\n✓ Running redaction script on: {docx_path}")
    result = subprocess.run(
        [sys.executable, str(script_path), str(docx_path)],
        capture_output=True,
        text=True
    )

    print(result.stdout)
    if result.stderr:
        print("STDERR:", result.stderr, file=sys.stderr)

    return result.returncode


def verify_redaction(md_path: Path, student_info: dict) -> bool:
    """
    Verify that the markdown file has been properly redacted.
    Returns True if all student info is redacted, False otherwise.
    """
    print(f"\n✓ Verifying redacted output: {md_path}")

    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    print("\n--- Redacted Content ---")
    print(content)
    print("--- End Content ---\n")

    # Check that student information is NOT present in the redacted content
    failures = []

    if student_info['name'] in content:
        failures.append(f"Student name '{student_info['name']}' found in redacted content")

    if student_info['email'] in content:
        failures.append(f"Student email '{student_info['email']}' found in redacted content")

    if student_info['student_id'] in content:
        failures.append(f"Student ID '{student_info['student_id']}' found in redacted content")

    # Check that redaction placeholders ARE present
    expected_placeholders = ['[STUDENT_NAME]', '[STUDENT_EMAIL]', '[STUDENT_ID]']
    placeholder_found = False
    for placeholder in expected_placeholders:
        if placeholder in content:
            placeholder_found = True
            print(f"✓ Found expected placeholder: {placeholder}")

    if not placeholder_found:
        failures.append("No redaction placeholders ([STUDENT_NAME], [STUDENT_EMAIL], [STUDENT_ID]) found")

    # Check that professor information is still present (should NOT be redacted)
    if 'Professor Jane Doe' not in content:
        failures.append("Professor name was incorrectly redacted")

    if 'Department of Computer Science' not in content:
        failures.append("Department information was incorrectly redacted")

    # Report results
    if failures:
        print("\n✗ REDACTION VERIFICATION FAILED:")
        for failure in failures:
            print(f"  - {failure}")
        return False
    else:
        print("\n✓ REDACTION VERIFICATION PASSED")
        print("  - All student information properly redacted")
        print("  - Redaction placeholders present")
        print("  - Non-student information preserved")
        return True


def main():
    """Run the system test."""
    print("="*60)
    print("Student Information Redaction - System Test")
    print("="*60)

    # Check for OpenAI API key
    if not os.getenv('OPENAI_API_KEY'):
        print("\n✗ ERROR: OPENAI_API_KEY environment variable not set")
        print("Please set your OpenAI API key before running tests")
        sys.exit(1)

    # Create a temporary directory for test files
    with tempfile.TemporaryDirectory() as temp_dir:
        temp_path = Path(temp_dir)
        docx_path = temp_path / 'test_letter.docx'
        md_path = temp_path / 'test_letter.md'

        # Step 1: Create test document
        print("\n[1/3] Creating test document...")
        student_info = create_test_document(docx_path)

        # Step 2: Run redaction script
        print("\n[2/3] Running redaction script...")
        exit_code = run_redaction_script(docx_path)

        if exit_code != 0:
            print(f"\n✗ TEST FAILED: Redaction script exited with code {exit_code}")
            sys.exit(1)

        # Check that output file was created
        if not md_path.exists():
            print(f"\n✗ TEST FAILED: Expected output file not created: {md_path}")
            sys.exit(1)

        # Step 3: Verify redaction
        print("\n[3/3] Verifying redaction...")
        if verify_redaction(md_path, student_info):
            print("\n" + "="*60)
            print("✓ ALL TESTS PASSED")
            print("="*60)
            sys.exit(0)
        else:
            print("\n" + "="*60)
            print("✗ TESTS FAILED")
            print("="*60)
            sys.exit(1)


if __name__ == '__main__':
    main()
