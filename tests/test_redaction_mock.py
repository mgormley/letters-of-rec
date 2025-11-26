#!/usr/bin/env python3
"""
System test for the redaction script (without requiring OpenAI API).
Creates a sample .docx file, verifies conversion to markdown works,
and tests the script's basic functionality without making API calls.
"""

import subprocess
import sys
import tempfile
from pathlib import Path

from docx import Document


def create_test_document(output_path: Path) -> dict:
    """
    Create a test Word document with known student information.
    Returns a dict of student info that should be redacted.
    """
    doc = Document()

    # Add title
    doc.add_heading('Letter of Recommendation', 0)

    # Add content with student information
    student_info = {
        'name': 'John Smith',
        'email': 'jsmith@university.edu',
        'student_id': 'STU123456'
    }

    doc.add_paragraph(
        f"To Whom It May Concern,\n\n"
        f"I am writing to recommend {student_info['name']} for your graduate program. "
        f"{student_info['name']} has been an outstanding student in my Computer Science course.\n\n"
        f"You can reach {student_info['name']} at {student_info['email']} or reference "
        f"student ID {student_info['student_id']} in your records.\n\n"
        f"{student_info['name']} has demonstrated excellent problem-solving skills and "
        f"a deep understanding of algorithms and data structures.\n\n"
        f"Sincerely,\n"
        f"Professor Jane Doe\n"
        f"Department of Computer Science\n"
        f"University of Example"
    )

    doc.save(str(output_path))
    print(f"✓ Created test document: {output_path}")
    return student_info


def test_dry_run_cli(docx_path: Path) -> bool:
    """
    Test the CLI in dry-run mode (no API calls).
    Returns True if successful, False otherwise.
    """
    print(f"\n✓ Running CLI redact command in dry-run mode on: {docx_path}")
    result = subprocess.run(
        [sys.executable, '-m', 'lor.cli', 'redact', str(docx_path), '--dry-run'],
        capture_output=True,
        text=True
    )

    print(result.stdout)
    if result.stderr:
        print("STDERR:", result.stderr, file=sys.stderr)

    return result.returncode == 0


def test_dry_run_script(docx_path: Path) -> bool:
    """
    Test the standalone script in dry-run mode (no API calls).
    Returns True if successful, False otherwise.
    """
    print(f"\n✓ Running standalone script in dry-run mode on: {docx_path}")
    result = subprocess.run(
        [sys.executable, '-m', 'lor.redact_student_info', str(docx_path), '--dry-run'],
        capture_output=True,
        text=True
    )

    print(result.stdout)
    if result.stderr:
        print("STDERR:", result.stderr, file=sys.stderr)

    return result.returncode == 0


def test_conversion(docx_path: Path) -> bool:
    """
    Test that we can convert the docx to markdown using the mammoth library.
    """
    import mammoth

    print(f"\n✓ Testing Word to Markdown conversion...")

    try:
        with open(docx_path, 'rb') as docx_file:
            result = mammoth.convert_to_markdown(docx_file)
            markdown_text = result.value

        if not markdown_text:
            print("✗ Conversion produced empty output")
            return False

        print(f"✓ Conversion successful ({len(markdown_text)} characters)")
        print("\n--- Converted Markdown ---")
        print(markdown_text)
        print("--- End Markdown ---\n")

        # Check that key content is present
        if "Letter of Recommendation" not in markdown_text:
            print("✗ Title not found in converted markdown")
            return False

        if "John Smith" not in markdown_text:
            print("✗ Student name not found in converted markdown")
            return False

        if "Professor Jane Doe" not in markdown_text:
            print("✗ Professor name not found in converted markdown")
            return False

        print("✓ All expected content found in markdown")
        return True

    except Exception as e:
        print(f"✗ Conversion failed: {e}")
        return False


def main():
    """Run the system test."""
    print("="*60)
    print("Student Information Redaction - System Test (Mock)")
    print("="*60)
    print("\nNote: This test runs without making OpenAI API calls")
    print("It validates document creation, conversion, and script execution\n")

    # Create a temporary directory for test files
    with tempfile.TemporaryDirectory() as temp_dir:
        temp_path = Path(temp_dir)
        docx_path = temp_path / 'test_letter.docx'

        # Step 1: Create test document
        print("\n[1/3] Creating test document...")
        student_info = create_test_document(docx_path)

        # Verify the file exists
        if not docx_path.exists():
            print(f"✗ TEST FAILED: Document not created at {docx_path}")
            sys.exit(1)

        print(f"✓ Document created successfully ({docx_path.stat().st_size} bytes)")

        # Step 2: Test conversion
        print("\n[2/3] Testing Word to Markdown conversion...")
        if not test_conversion(docx_path):
            print("\n✗ TEST FAILED: Conversion test failed")
            sys.exit(1)

        # Step 3: Test CLI in dry-run mode
        print("\n[3/4] Testing CLI execution (dry-run mode)...")
        if not test_dry_run_cli(docx_path):
            print("\n✗ TEST FAILED: CLI execution failed")
            sys.exit(1)

        # Step 4: Test standalone script in dry-run mode
        print("\n[4/4] Testing standalone script execution (dry-run mode)...")
        if not test_dry_run_script(docx_path):
            print("\n✗ TEST FAILED: Script execution failed")
            sys.exit(1)

        print("\n" + "="*60)
        print("✓ ALL TESTS PASSED")
        print("="*60)
        print("\nBoth CLI and standalone script are ready to use!")
        print("To test with actual redaction, set OPENAI_API_KEY and run:")
        print("  python3 tests/test_redaction.py")


if __name__ == '__main__':
    main()
